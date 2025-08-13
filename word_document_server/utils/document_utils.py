"""
Document utility functions for Word Document Server.
"""
import json
from typing import Dict, List, Any, Optional, Tuple
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml.ns import qn


def get_document_properties(doc_path: str) -> Dict[str, Any]:
    """Get properties of a Word document."""
    import os
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}
    
    try:
        doc = Document(doc_path)
        core_props = doc.core_properties
        
        return {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "keywords": core_props.keywords or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
            "last_modified_by": core_props.last_modified_by or "",
            "revision": core_props.revision or 0,
            "page_count": len(doc.sections),
            "word_count": sum(len(paragraph.text.split()) for paragraph in doc.paragraphs),
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables)
        }
    except Exception as e:
        return {"error": f"Failed to get document properties: {str(e)}"}


def extract_document_text(doc_path: str, max_chars: Optional[int] = None, 
                         include_tables: bool = True) -> str:
    """
    Extract all text from a Word document with optional size limits.
    
    Args:
        doc_path: Path to the Word document
        max_chars: Maximum characters to return (None for no limit)
        include_tables: Whether to include table content
    """
    import os
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    
    try:
        doc = Document(doc_path)
        text_parts = []
        total_chars = 0
        
        # Extract paragraph text
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
                total_chars += len(paragraph.text)
                if max_chars and total_chars > max_chars:
                    break
        
        # Extract table text if requested
        if include_tables:
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_content = " ".join(p.text for p in cell.paragraphs if p.text.strip())
                        if cell_content:
                            row_text.append(cell_content)
                    if row_text:
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    text_parts.append("\n[Table]\n" + "\n".join(table_text))
                    total_chars += len(text_parts[-1])
                    if max_chars and total_chars > max_chars:
                        break
        
        result = "\n\n".join(text_parts)
        
        # Apply character limit if specified
        if max_chars and len(result) > max_chars:
            result = result[:max_chars] + "... [content truncated]"
        
        return result
        
    except Exception as e:
        return f"Failed to extract text: {str(e)}"


def extract_document_text_chunked(doc_path: str, chunk_size: int = 50000) -> Tuple[str, bool]:
    """
    Extract document text in chunks for large files.
    
    Args:
        doc_path: Path to the Word document
        chunk_size: Maximum characters per chunk
        
    Returns:
        Tuple of (text_content, is_truncated)
    """
    import os
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist", False
    
    try:
        doc = Document(doc_path)
        text_parts = []
        total_chars = 0
        is_truncated = False
        
        # Extract paragraph text
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                if total_chars + len(paragraph.text) > chunk_size:
                    remaining = chunk_size - total_chars
                    if remaining > 0:
                        text_parts.append(paragraph.text[:remaining])
                    is_truncated = True
                    break
                
                text_parts.append(paragraph.text)
                total_chars += len(paragraph.text)
        
        # Add table content if space permits
        if not is_truncated:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = " ".join(p.text for p in cell.paragraphs if p.text.strip())
                        if cell_text:
                            if total_chars + len(cell_text) > chunk_size:
                                remaining = chunk_size - total_chars
                                if remaining > 0:
                                    text_parts.append(cell_text[:remaining])
                                is_truncated = True
                                break
                            text_parts.append(cell_text)
                            total_chars += len(cell_text)
                    if is_truncated:
                        break
                if is_truncated:
                    break
        
        result = "\n\n".join(text_parts)
        return result, is_truncated
        
    except Exception as e:
        return f"Failed to extract text: {str(e)}", False


def get_document_structure(doc_path: str) -> Dict[str, Any]:
    """Get the structure of a Word document."""
    import os
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}
    
    try:
        doc = Document(doc_path)
        structure = {
            "paragraphs": [],
            "tables": []
        }
        
        # Get paragraphs
        for i, para in enumerate(doc.paragraphs):
            structure["paragraphs"].append({
                "index": i,
                "text": para.text[:100] + ("..." if len(para.text) > 100 else ""),
                "style": para.style.name if para.style else "Normal"
            })
        
        # Get tables
        for i, table in enumerate(doc.tables):
            table_data = {
                "index": i,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "preview": []
            }
            
            # Get sample of table data
            max_rows = min(3, len(table.rows))
            for row_idx in range(max_rows):
                row_data = []
                max_cols = min(3, len(table.columns))
                for col_idx in range(max_cols):
                    try:
                        cell_text = table.cell(row_idx, col_idx).text
                        row_data.append(cell_text[:20] + ("..." if len(cell_text) > 20 else ""))
                    except IndexError:
                        row_data.append("N/A")
                table_data["preview"].append(row_data)
            
            structure["tables"].append(table_data)
        
        return structure
    except Exception as e:
        return {"error": f"Failed to get document structure: {str(e)}"}


def get_all_paragraphs(doc_path: str) -> Dict[str, Any]:
    """
    一次性获取所有段落内容
    
    Args:
        doc_path: Word文档路径
    
    Returns:
        包含所有段落信息的字典
    """
    import os
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}
    
    try:
        doc = Document(doc_path)
        paragraphs = []
        
        for i, para in enumerate(doc.paragraphs):
            paragraphs.append({
                "index": i,
                "text": para.text,
                "style": para.style.name if para.style else "Normal",
                "runs": [
                    {
                        "text": run.text,
                        "bold": run.bold,
                        "italic": run.italic,
                        "underline": run.underline
                    }
                    for run in para.runs if run.text
                ]
            })
        
        return {
            "total_paragraphs": len(paragraphs),
            "paragraphs": paragraphs
        }
    except Exception as e:
        return {"error": f"Failed to get paragraphs: {str(e)}"}


def get_paragraphs_by_range(doc_path: str, start_index: int = 0, end_index: Optional[int] = None) -> Dict[str, Any]:
    """
    获取指定段落范围的内容
    
    Args:
        doc_path: Word文档路径
        start_index: 起始段落索引（包含）
        end_index: 结束段落索引（不包含），None表示到最后
    
    Returns:
        包含指定范围段落信息的字典
    """
    import os
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}
    
    try:
        doc = Document(doc_path)
        paragraphs = []
        
        total_paragraphs = len(doc.paragraphs)
        
        # 验证索引范围
        if start_index < 0:
            start_index = 0
        if end_index is None:
            end_index = total_paragraphs
        elif end_index > total_paragraphs:
            end_index = total_paragraphs
        
        if start_index >= total_paragraphs:
            return {
                "error": f"Start index {start_index} exceeds total paragraphs {total_paragraphs}"
            }
        
        # 获取指定范围的段落
        for i in range(start_index, min(end_index, total_paragraphs)):
            para = doc.paragraphs[i]
            paragraphs.append({
                "index": i,
                "text": para.text,
                "style": para.style.name if para.style else "Normal",
                "runs": [
                    {
                        "text": run.text,
                        "bold": run.bold,
                        "italic": run.italic,
                        "underline": run.underline,
                        "font_size": run.font.size.pt if run.font.size else None
                    }
                    for run in para.runs if run.text
                ]
            })
        
        return {
            "total_paragraphs": total_paragraphs,
            "range": f"{start_index}-{min(end_index, total_paragraphs)}",
            "paragraphs": paragraphs
        }
    except Exception as e:
        return {"error": f"Failed to get paragraphs by range: {str(e)}"}


def get_paragraphs_by_page(doc_path: str, page_number: int, page_size: int = 100) -> Dict[str, Any]:
    """
    分页获取段落内容
    
    Args:
        doc_path: Word文档路径
        page_number: 页码（从1开始）
        page_size: 每页段落数量
    
    Returns:
        包含分页段落信息的字典
    """
    import os
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}
    
    try:
        doc = Document(doc_path)
        total_paragraphs = len(doc.paragraphs)
        
        # 计算分页信息
        total_pages = (total_paragraphs + page_size - 1) // page_size
        
        if page_number < 1:
            page_number = 1
        if page_number > total_pages:
            page_number = total_pages
        
        start_index = (page_number - 1) * page_size
        end_index = min(start_index + page_size, total_paragraphs)
        
        paragraphs = []
        for i in range(start_index, end_index):
            para = doc.paragraphs[i]
            paragraphs.append({
                "index": i,
                "text": para.text,
                "style": para.style.name if para.style else "Normal",
                "character_count": len(para.text),
                "word_count": len(para.text.split())
            })
        
        return {
            "total_paragraphs": total_paragraphs,
            "total_pages": total_pages,
            "current_page": page_number,
            "page_size": page_size,
            "range": f"{start_index}-{end_index-1}",
            "paragraphs": paragraphs
        }
    except Exception as e:
        return {"error": f"Failed to get paragraphs by page: {str(e)}"}


def analyze_paragraph_distribution(doc_path: str) -> Dict[str, Any]:
    """
    分析段落分布情况
    
    Args:
        doc_path: Word文档路径
    
    Returns:
        段落统计分析信息
    """
    import os
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}
    
    try:
        doc = Document(doc_path)
        
        if not doc.paragraphs:
            return {
                "total_paragraphs": 0,
                "empty_paragraphs": 0,
                "non_empty_paragraphs": 0,
                "styles": {},
                "character_stats": {
                    "min": 0,
                    "max": 0,
                    "avg": 0,
                    "total": 0
                }
            }
        
        paragraphs = doc.paragraphs
        total_paragraphs = len(paragraphs)
        
        # 统计分析
        empty_paragraphs = sum(1 for p in paragraphs if not p.text.strip())
        non_empty_paragraphs = total_paragraphs - empty_paragraphs
        
        # 样式统计
        styles = {}
        for para in paragraphs:
            style_name = para.style.name if para.style else "Normal"
            styles[style_name] = styles.get(style_name, 0) + 1
        
        # 字符统计
        char_counts = [len(p.text) for p in paragraphs if p.text.strip()]
        if char_counts:
            char_stats = {
                "min": min(char_counts),
                "max": max(char_counts),
                "avg": sum(char_counts) // len(char_counts),
                "total": sum(char_counts)
            }
        else:
            char_stats = {"min": 0, "max": 0, "avg": 0, "total": 0}
        
        return {
            "total_paragraphs": total_paragraphs,
            "empty_paragraphs": empty_paragraphs,
            "non_empty_paragraphs": non_empty_paragraphs,
            "styles": styles,
            "character_stats": char_stats,
            "word_stats": {
                "total_words": sum(len(p.text.split()) for p in paragraphs),
                "avg_words_per_paragraph": sum(len(p.text.split()) for p in paragraphs) // total_paragraphs if total_paragraphs > 0 else 0
            }
        }
    except Exception as e:
        return {"error": f"Failed to analyze paragraph distribution: {str(e)}"}


def find_paragraph_by_text(doc, text, partial_match=False):
    """
    Find paragraphs containing specific text.
    
    Args:
        doc: Document object
        text: Text to search for
        partial_match: If True, matches paragraphs containing the text; if False, matches exact text
        
    Returns:
        List of paragraph indices that match the criteria
    """
    matching_paragraphs = []
    
    for i, para in enumerate(doc.paragraphs):
        if partial_match and text in para.text:
            matching_paragraphs.append(i)
        elif not partial_match and para.text == text:
            matching_paragraphs.append(i)
            
    return matching_paragraphs


def find_and_replace_text(doc, old_text, new_text):
    """
    Find and replace text throughout the document, skipping Table of Contents (TOC) paragraphs.
    
    Args:
        doc: Document object
        old_text: Text to find
        new_text: Text to replace with
        
    Returns:
        Number of replacements made
    """
    count = 0
    
    # Search in paragraphs
    for para in doc.paragraphs:
        # Skip TOC paragraphs
        if para.style and para.style.name.startswith("TOC"):
            continue
        if old_text in para.text:
            for run in para.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    count += 1
    
    # Search in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    # Skip TOC paragraphs in tables
                    if para.style and para.style.name.startswith("TOC"):
                        continue
                    if old_text in para.text:
                        for run in para.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
                                count += 1
    
    return count


def get_document_xml(doc_path: str) -> str:
    """Extract and return the raw XML structure of the Word document (word/document.xml)."""
    import os
    import zipfile
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    try:
        with zipfile.ZipFile(doc_path) as docx_zip:
            with docx_zip.open('word/document.xml') as xml_file:
                return xml_file.read().decode('utf-8')
    except Exception as e:
        return f"Failed to extract XML: {str(e)}"


def insert_header_near_text(doc_path: str, target_text: str = None, header_title: str = "", position: str = 'after', header_style: str = 'Heading 1', target_paragraph_index: int = None) -> str:
    """Insert a header (with specified style) before or after the target paragraph. Specify by text or paragraph index. Skips TOC paragraphs in text search."""
    import os
    from docx import Document
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    try:
        doc = Document(doc_path)
        found = False
        para = None
        if target_paragraph_index is not None:
            if target_paragraph_index < 0 or target_paragraph_index >= len(doc.paragraphs):
                return f"Invalid target_paragraph_index: {target_paragraph_index}. Document has {len(doc.paragraphs)} paragraphs."
            para = doc.paragraphs[target_paragraph_index]
            found = True
        else:
            for i, p in enumerate(doc.paragraphs):
                # Skip TOC paragraphs
                if p.style and p.style.name.lower().startswith("toc"):
                    continue
                if target_text and target_text in p.text:
                    para = p
                    found = True
                    break
        if not found or para is None:
            return f"Target paragraph not found (by index or text). (TOC paragraphs are skipped in text search)"
        # Save anchor index before insertion
        if target_paragraph_index is not None:
            anchor_index = target_paragraph_index
        else:
            anchor_index = None
            for i, p in enumerate(doc.paragraphs):
                if p is para:
                    anchor_index = i
                    break
        new_para = doc.add_paragraph(header_title, style=header_style)
        if position == 'before':
            para._element.addprevious(new_para._element)
        else:
            para._element.addnext(new_para._element)
        doc.save(doc_path)
        if anchor_index is not None:
            return f"Header '{header_title}' (style: {header_style}) inserted {position} paragraph (index {anchor_index})."
        else:
            return f"Header '{header_title}' (style: {header_style}) inserted {position} the target paragraph."
    except Exception as e:
        return f"Failed to insert header: {str(e)}"


def insert_line_or_paragraph_near_text(doc_path: str, target_text: str = None, line_text: str = "", position: str = 'after', line_style: str = None, target_paragraph_index: int = None) -> str:
    """
    Insert a new line or paragraph (with specified or matched style) before or after the target paragraph.
    You can specify the target by text (first match) or by paragraph index.
    Skips paragraphs whose style name starts with 'TOC' if using text search.
    """
    import os
    from docx import Document
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    try:
        doc = Document(doc_path)
        found = False
        para = None
        if target_paragraph_index is not None:
            if target_paragraph_index < 0 or target_paragraph_index >= len(doc.paragraphs):
                return f"Invalid target_paragraph_index: {target_paragraph_index}. Document has {len(doc.paragraphs)} paragraphs."
            para = doc.paragraphs[target_paragraph_index]
            found = True
        else:
            for i, p in enumerate(doc.paragraphs):
                # Skip TOC paragraphs
                if p.style and p.style.name.lower().startswith("toc"):
                    continue
                if target_text and target_text in p.text:
                    para = p
                    found = True
                    break
        if not found or para is None:
            return f"Target paragraph not found (by index or text). (TOC paragraphs are skipped in text search)"
        # Save anchor index before insertion
        if target_paragraph_index is not None:
            anchor_index = target_paragraph_index
        else:
            anchor_index = None
            for i, p in enumerate(doc.paragraphs):
                if p is para:
                    anchor_index = i
                    break
        # Determine style: use provided or match target
        style = line_style if line_style else para.style
        new_para = doc.add_paragraph(line_text, style=style)
        if position == 'before':
            para._element.addprevious(new_para._element)
        else:
            para._element.addnext(new_para._element)
        doc.save(doc_path)
        if anchor_index is not None:
            return f"Line/paragraph inserted {position} paragraph (index {anchor_index}) with style '{style}'."
        else:
            return f"Line/paragraph inserted {position} the target paragraph with style '{style}'."
    except Exception as e:
        return f"Failed to insert line/paragraph: {str(e)}"


def insert_numbered_list_near_text(doc_path: str, target_text: str = None, list_items: list = None, position: str = 'after', target_paragraph_index: int = None) -> str:
    """
    Insert a numbered list before or after the target paragraph. Specify by text or paragraph index. Skips TOC paragraphs in text search.
    Args:
        doc_path: Path to the Word document
        target_text: Text to search for in paragraphs (optional if using index)
        list_items: List of strings, each as a list item
        position: 'before' or 'after' (default: 'after')
        target_paragraph_index: Optional paragraph index to use as anchor
    Returns:
        Status message
    """
    import os
    from docx import Document
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    try:
        doc = Document(doc_path)
        found = False
        para = None
        if target_paragraph_index is not None:
            if target_paragraph_index < 0 or target_paragraph_index >= len(doc.paragraphs):
                return f"Invalid target_paragraph_index: {target_paragraph_index}. Document has {len(doc.paragraphs)} paragraphs."
            para = doc.paragraphs[target_paragraph_index]
            found = True
        else:
            for i, p in enumerate(doc.paragraphs):
                # Skip TOC paragraphs
                if p.style and p.style.name.lower().startswith("toc"):
                    continue
                if target_text and target_text in p.text:
                    para = p
                    found = True
                    break
        if not found or para is None:
            return f"Target paragraph not found (by index or text). (TOC paragraphs are skipped in text search)"
        # Save anchor index before insertion
        if target_paragraph_index is not None:
            anchor_index = target_paragraph_index
        else:
            anchor_index = None
            for i, p in enumerate(doc.paragraphs):
                if p is para:
                    anchor_index = i
                    break
        # Robust style selection for numbered list
        style_name = None
        for candidate in ['List Number', 'List Paragraph', 'Normal']:
            try:
                _ = doc.styles[candidate]
                style_name = candidate
                break
            except KeyError:
                continue
        if not style_name:
            style_name = None  # fallback to default
        new_paras = []
        for item in (list_items or []):
            p = doc.add_paragraph(item, style=style_name)
            new_paras.append(p)
        # Move the new paragraphs to the correct position
        for p in reversed(new_paras):
            if position == 'before':
                para._element.addprevious(p._element)
            else:
                para._element.addnext(p._element)
        doc.save(doc_path)
        if anchor_index is not None:
            return f"Numbered list inserted {position} paragraph (index {anchor_index})."
        else:
            return f"Numbered list inserted {position} the target paragraph."
    except Exception as e:
        return f"Failed to insert numbered list: {str(e)}"


def is_toc_paragraph(para):
    """Devuelve True si el párrafo tiene un estilo de tabla de contenido (TOC)."""
    return para.style and para.style.name.upper().startswith("TOC")


def is_heading_paragraph(para):
    """Devuelve True si el párrafo tiene un estilo de encabezado (Heading 1, Heading 2, etc)."""
    return para.style and para.style.name.lower().startswith("heading")


# --- Helper: Get style name from a <w:p> element ---
def get_paragraph_style(el):
    from docx.oxml.ns import qn
    pPr = el.find(qn('w:pPr'))
    if pPr is not None:
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is not None and 'w:val' in pStyle.attrib:
            return pStyle.attrib['w:val']
    return None

# --- Main: Delete everything under a header until next heading/TOC ---
def delete_block_under_header(doc, header_text):
    """
    Remove all elements (paragraphs, tables, etc.) after the header (by text) and before the next heading/TOC (by style).
    Returns: (header_element, elements_removed)
    """
    # Find the header paragraph by text (like delete_paragraph finds by index)
    header_para = None
    header_idx = None
    
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().lower() == header_text.strip().lower():
            header_para = para
            header_idx = i
            break
    
    if header_para is None:
        return None, 0
    
    # Find the next heading/TOC paragraph to determine the end of the block
    end_idx = None
    for i in range(header_idx + 1, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        if para.style and para.style.name.lower().startswith(('heading', 'título', 'toc')):
            end_idx = i
            break
    
    # If no next heading found, delete until end of document
    if end_idx is None:
        end_idx = len(doc.paragraphs)
    
    # Remove paragraphs by index (like delete_paragraph does)
    removed_count = 0
    for i in range(header_idx + 1, end_idx):
        if i < len(doc.paragraphs):  # Safety check
            para = doc.paragraphs[header_idx + 1]  # Always remove the first paragraph after header
            p = para._p
            p.getparent().remove(p)
            removed_count += 1
    
    return header_para._p, removed_count

# --- Usage in replace_paragraph_block_below_header ---
def replace_paragraph_block_below_header(
    doc_path: str,
    header_text: str,
    new_paragraphs: list,
    detect_block_end_fn=None,
    new_paragraph_style: str = None
) -> str:
    """
    Reemplaza todo el contenido debajo de una cabecera (por texto), hasta el siguiente encabezado/TOC (por estilo).
    """
    from docx import Document
    import os
    if not os.path.exists(doc_path):
        return f"Document {doc_path} not found."
    
    doc = Document(doc_path)
    
    # Find the header paragraph first
    header_para = None
    header_idx = None
    for i, para in enumerate(doc.paragraphs):
        para_text = para.text.strip().lower()
        is_toc = is_toc_paragraph(para)
        if para_text == header_text.strip().lower() and not is_toc:
            header_para = para
            header_idx = i
            break
    
    if header_para is None:
        return f"Header '{header_text}' not found in document."
    
    # Delete everything under the header using the same document instance
    header_el, removed_count = delete_block_under_header(doc, header_text)
    
    # Now insert new paragraphs after the header (which should still be in the document)
    style_to_use = new_paragraph_style or "Normal"
    
    # Find the header again after deletion (it should still be there)
    current_para = header_para
    for text in new_paragraphs:
        new_para = doc.add_paragraph(text, style=style_to_use)
        current_para._element.addnext(new_para._element)
        current_para = new_para
    
    doc.save(doc_path)
    return f"Replaced content under '{header_text}' with {len(new_paragraphs)} paragraph(s), style: {style_to_use}, removed {removed_count} elements."


def replace_block_between_manual_anchors(
    doc_path: str,
    start_anchor_text: str,
    new_paragraphs: list,
    end_anchor_text: str = None,
    match_fn=None,
    new_paragraph_style: str = None
) -> str:
    """
    Replace all content (paragraphs, tables, etc.) between start_anchor_text and end_anchor_text (or next logical header if not provided).
    If end_anchor_text is None, deletes until next visually distinct paragraph (bold, all caps, or different font size), or end of document.
    Inserts new_paragraphs after the start anchor.
    """
    from docx import Document
    import os
    if not os.path.exists(doc_path):
        return f"Document {doc_path} not found."
    doc = Document(doc_path)
    body = doc.element.body
    elements = list(body)
    start_idx = None
    end_idx = None
    # Find start anchor
    for i, el in enumerate(elements):
        if el.tag == CT_P.tag:
            p_text = "".join([node.text or '' for node in el.iter() if node.tag.endswith('}t')]).strip()
            if match_fn:
                if match_fn(p_text, el):
                    start_idx = i
                    break
            elif p_text == start_anchor_text.strip():
                start_idx = i
                break
    if start_idx is None:
        return f"Start anchor '{start_anchor_text}' not found."
    # Find end anchor
    if end_anchor_text:
        for i in range(start_idx + 1, len(elements)):
            el = elements[i]
            if el.tag == CT_P.tag:
                p_text = "".join([node.text or '' for node in el.iter() if node.tag.endswith('}t')]).strip()
                if match_fn:
                    if match_fn(p_text, el, is_end=True):
                        end_idx = i
                        break
                elif p_text == end_anchor_text.strip():
                    end_idx = i
                    break
    else:
        # Heuristic: next visually distinct paragraph (bold, all caps, or different font size), or end of document
        for i in range(start_idx + 1, len(elements)):
            el = elements[i]
            if el.tag == CT_P.tag:
                # Check for bold, all caps, or font size
                runs = [node for node in el.iter() if node.tag.endswith('}r')]
                for run in runs:
                    rpr = run.find(qn('w:rPr'))
                    if rpr is not None:
                        if rpr.find(qn('w:b')) is not None or rpr.find(qn('w:caps')) is not None or rpr.find(qn('w:sz')) is not None:
                            end_idx = i
                            break
                if end_idx is not None:
                    break
    # Mark elements for removal
    to_remove = []
    for i in range(start_idx + 1, end_idx if end_idx is not None else len(elements)):
        to_remove.append(elements[i])
    for el in to_remove:
        body.remove(el)
    doc.save(doc_path)
    # Reload and find start anchor for insertion
    doc = Document(doc_path)
    paras = doc.paragraphs
    anchor_idx = None
    for i, para in enumerate(paras):
        if para.text.strip() == start_anchor_text.strip():
            anchor_idx = i
            break
    if anchor_idx is None:
        return f"Start anchor '{start_anchor_text}' not found after deletion (unexpected)."
    anchor_para = paras[anchor_idx]
    style_to_use = new_paragraph_style or "Normal"
    for text in new_paragraphs:
        new_para = doc.add_paragraph(text, style=style_to_use)
        anchor_para._element.addnext(new_para._element)
        anchor_para = new_para
    doc.save(doc_path)
    return f"Replaced content between '{start_anchor_text}' and '{end_anchor_text or 'next logical header'}' with {len(new_paragraphs)} paragraph(s), style: {style_to_use}, removed {len(to_remove)} elements."
