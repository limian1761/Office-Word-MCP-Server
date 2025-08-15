from docx import Document

# Create a new Document
doc = Document()

doc.add_heading('测试文档', 1)

doc.add_paragraph('这是一个测试文档。')
doc.add_paragraph('文档中包含一些文本，用于测试查找和替换功能。')
doc.add_paragraph('测试文档的格式需要保持一致。')
doc.add_paragraph('这是文档的最后一行。')

doc.save('test_document.docx')
print('Test document created successfully.')